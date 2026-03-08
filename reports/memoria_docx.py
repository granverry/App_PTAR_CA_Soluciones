from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os


def _agregar_titulo(doc, texto, nivel=1):
    p = doc.add_paragraph()
    p.style = f"Heading {nivel}"
    run = p.add_run(texto)
    run.bold = True
    return p


def _agregar_parrafo(doc, texto, negrita=False, centrado=False):
    p = doc.add_paragraph()
    if centrado:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto)
    run.bold = negrita
    return p


def _agregar_tabla_diccionario(doc, diccionario, titulo=None):
    if titulo:
        _agregar_titulo(doc, titulo, nivel=3)

    tabla = doc.add_table(rows=1, cols=2)
    tabla.style = "Table Grid"

    hdr = tabla.rows[0].cells
    hdr[0].text = "Parámetro"
    hdr[1].text = "Valor"

    for k, v in diccionario.items():
        row = tabla.add_row().cells
        row[0].text = str(k)
        row[1].text = str(v)

    doc.add_paragraph("")


def _insertar_imagen(doc, ruta_imagen, titulo=None, ancho=6.3):
    if ruta_imagen and os.path.exists(ruta_imagen):
        if titulo:
            _agregar_titulo(doc, titulo, nivel=3)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(ruta_imagen, width=Inches(ancho))
        doc.add_paragraph("")


def _agregar_salto_pagina(doc):
    doc.add_page_break()


def _agregar_portada(doc, datos_proyecto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run("\n\nMEMORIA DE CÁLCULO\n")
    run.bold = True
    run.font.size = Pt(18)

    run2 = p.add_run("PLANTA DE TRATAMIENTO DE AGUAS RESIDUALES\n\n")
    run2.bold = True
    run2.font.size = Pt(16)

    nombre_proyecto = datos_proyecto.get("Nombre del proyecto", "Proyecto PTAR")
    run3 = p.add_run(f"{nombre_proyecto}\n\n")
    run3.font.size = Pt(14)

    ubicacion = datos_proyecto.get("Ubicación", "No especificada")
    p.add_run(f"Ubicación: {ubicacion}\n")

    poblacion = datos_proyecto.get("Población", "No especificada")
    p.add_run(f"Población servida: {poblacion}\n")

    caudal = datos_proyecto.get("Caudal de entrada (L/s)", "No especificado")
    p.add_run(f"Caudal de diseño: {caudal} L/s\n")

    temperatura = datos_proyecto.get("Temperatura de diseño (°C)", "No especificada")
    p.add_run(f"Temperatura de diseño: {temperatura} °C\n")

    p.add_run("\nDesarrollado con PTAR Designer\n")

    _agregar_salto_pagina(doc)


def _resumen_ejecutivo(resultados):
    resumen = {}

    claves = [
        "Volumen del reactor (m3)",
        "TRH recalculado (h)",
        "SSVLM (mg/L)",
        "Relación F/M calculada",
        "Caudal de recirculación (m3/d)",
        "Caudal de purga de lodos (m3/d)",
        "Requerimiento de oxígeno (kg O2/d)",
        "Caudal de aire corregido (Nm3/h)",
        "Potencia requerida (kW)",
        "Número de difusores"
    ]

    for clave in claves:
        if clave in resultados:
            resumen[clave] = resultados[clave]

    return resumen


def _separar_resultados(resultados):
    bloque1 = {
        "DBO soluble que escapa del tratamiento (mg/L)": resultados.get("DBO soluble que escapa del tratamiento (mg/L)"),
        "Eficiencia tratamiento biológico": resultados.get("Eficiencia tratamiento biológico"),
        "Eficiencia conjunta de planta": resultados.get("Eficiencia conjunta de planta"),
    }

    bloque2 = {
        "Tiempo de retención celular mínima (d)": resultados.get("Tiempo de retención celular mínima (d)"),
        "DBO salida recalculada (mg/L)": resultados.get("DBO salida recalculada (mg/L)"),
        "SSVLM (mg/L)": resultados.get("SSVLM (mg/L)"),
        "Volumen del reactor (m3)": resultados.get("Volumen del reactor (m3)"),
        "Relación F/M calculada": resultados.get("Relación F/M calculada"),
        "TRH recalculado (h)": resultados.get("TRH recalculado (h)"),
        "Área tanque (m2)": resultados.get("Área tanque (m2)"),
        "Ancho tanque rectangular (m)": resultados.get("Ancho tanque rectangular (m)"),
        "Longitud tanque rectangular (m)": resultados.get("Longitud tanque rectangular (m)"),
    }

    bloque3 = {
        "Rendimiento observado": resultados.get("Rendimiento observado"),
        "Masa lodo volatil purgado (kg/d)": resultados.get("Masa lodo volatil purgado (kg/d)"),
        "Masa lodo total (kg/d)": resultados.get("Masa lodo total (kg/d)"),
        "Masa lodo a purgar (kg/d)": resultados.get("Masa lodo a purgar (kg/d)"),
        "Caudal de purga de lodos (m3/d)": resultados.get("Caudal de purga de lodos (m3/d)"),
        "Caudal de recirculación (m3/d)": resultados.get("Caudal de recirculación (m3/d)"),
        "Relación de recirculación": resultados.get("Relación de recirculación"),
    }

    bloque4 = {
        "Demanda por síntesis (kg O2/d)": resultados.get("Demanda por síntesis (kg O2/d)"),
        "Demanda endógena (kg O2/d)": resultados.get("Demanda endógena (kg O2/d)"),
        "Demanda por nitrificación (kg O2/d)": resultados.get("Demanda por nitrificación (kg O2/d)"),
        "Requerimiento de oxígeno (kg O2/d)": resultados.get("Requerimiento de oxígeno (kg O2/d)"),
        "Tasa transferencia estándar O2 (kg O2/h)": resultados.get("Tasa transferencia estándar O2 (kg O2/h)"),
        "Caudal de aire normal (Nm3/h)": resultados.get("Caudal de aire normal (Nm3/h)"),
        "Caudal de aire corregido (Nm3/h)": resultados.get("Caudal de aire corregido (Nm3/h)"),
        "Caudal de aire volumétrico (m3/h)": resultados.get("Caudal de aire volumétrico (m3/h)"),
        "Potencia requerida (kW)": resultados.get("Potencia requerida (kW)"),
        "Potencia requerida (HP)": resultados.get("Potencia requerida (HP)"),
        "Número de difusores": resultados.get("Número de difusores"),
    }

    return bloque1, bloque2, bloque3, bloque4


def _generar_observaciones(resultados):
    obs = []

    fm = resultados.get("Relación F/M calculada")
    trh = resultados.get("TRH recalculado (h)")
    ssvlm = resultados.get("SSVLM (mg/L)")
    aire = resultados.get("Caudal de aire corregido (Nm3/h)")

    if fm is not None:
        if fm < 0.05:
            obs.append("La relación F/M calculada es baja frente a valores típicos de operación.")
        elif fm > 0.25:
            obs.append("La relación F/M calculada es alta y debe revisarse el volumen o la biomasa de diseño.")
        else:
            obs.append("La relación F/M calculada se encuentra en un rango típico de diseño.")

    if trh is not None:
        if trh < 4:
            obs.append("El TRH calculado es bajo y puede comprometer la estabilidad del proceso.")
        elif trh > 24:
            obs.append("El TRH calculado es alto y puede indicar sobredimensionamiento.")
        else:
            obs.append("El TRH calculado se encuentra en un rango razonable para el proceso.")

    if ssvlm is not None:
        if ssvlm > 5000:
            obs.append("La concentración de SSVLM es alta y debe verificarse la capacidad de sedimentación secundaria.")
        else:
            obs.append("La concentración de SSVLM es compatible con un diseño convencional de lodos activados.")

    if aire is not None:
        obs.append(f"El caudal de aire corregido estimado para el sistema es de {aire} Nm³/h.")

    return obs


def generar_memoria(datos_proyecto, resultados, ruta_pfd_tren=None, ruta_pfd_lodos=None):
    doc = Document()

    # Márgenes
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    # Portada
    _agregar_portada(doc, datos_proyecto)

    # 1. Datos del proyecto
    _agregar_titulo(doc, "1. Datos del proyecto", nivel=1)
    _agregar_tabla_diccionario(doc, datos_proyecto)

    # 2. Resumen ejecutivo
    _agregar_titulo(doc, "2. Resumen ejecutivo", nivel=1)
    resumen = _resumen_ejecutivo(resultados)
    _agregar_tabla_diccionario(doc, resumen)

    # 3. Diagramas de proceso
    _agregar_titulo(doc, "3. Diagramas de proceso", nivel=1)
    _insertar_imagen(doc, ruta_pfd_tren, titulo="3.1 PFD general del tren de tratamiento")
    _insertar_imagen(doc, ruta_pfd_lodos, titulo="3.2 PFD del proceso de lodos activados")

    # 4. Resultados detallados
    _agregar_titulo(doc, "4. Resultados detallados del diseño", nivel=1)

    bloque1, bloque2, bloque3, bloque4 = _separar_resultados(resultados)

    _agregar_tabla_diccionario(doc, bloque1, titulo="4.1 Estimación de DBO en el efluente")
    _agregar_tabla_diccionario(doc, bloque2, titulo="4.2 Diseño del reactor biológico")
    _agregar_tabla_diccionario(doc, bloque3, titulo="4.3 Producción de lodos")
    _agregar_tabla_diccionario(doc, bloque4, titulo="4.4 Requerimiento de oxígeno y aireación")

    # 5. Observaciones
    _agregar_titulo(doc, "5. Observaciones de diseño", nivel=1)
    observaciones = _generar_observaciones(resultados)
    for obs in observaciones:
        doc.add_paragraph(obs, style="List Bullet")

    # 6. Conclusión
    _agregar_titulo(doc, "6. Conclusión", nivel=1)
    _agregar_parrafo(
        doc,
        "La presente memoria resume los principales resultados del dimensionamiento del sistema de lodos activados y su línea de aireación, con base en los datos de entrada y criterios adoptados en el aplicativo PTAR Designer."
    )

    ruta_salida = "memoria_calculo_ptar_v2.docx"
    doc.save(ruta_salida)

    return ruta_salida