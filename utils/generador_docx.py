from docx import Document

def generar_docx(resultados, titulo="Resumen"):
    doc = Document()
    doc.add_heading(f"{titulo}", level=1)
    for k, v in resultados.items():
        doc.add_paragraph(f"{k}: {v}")
    archivo = f"resumen_{titulo.lower().replace(' ', '_')}.docx"
    doc.save(archivo)
